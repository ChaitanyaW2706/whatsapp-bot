import csv
import io
import importlib
import pandas as pd
from docx import Document
from pypdf import PdfReader

Presentation = None
try:
    pptx_module = importlib.import_module('pptx')
    Presentation = pptx_module.Presentation
except ImportError:
    Presentation = None


def _paragraph_in_table(paragraph):
    """Return True if the paragraph belongs to a table cell."""
    element = paragraph._p
    while element is not None:
        if element.tag.endswith('}tbl'):
            return True
        element = element.getparent()
    return False


def process_excel(file_content):
    """
    Process Excel file and return a list of strings (sheet/row content).
    """
    excel_data = pd.read_excel(io.BytesIO(file_content), sheet_name=None)
    texts = []

    for sheet_name, df in excel_data.items():
        if df.empty:
            continue
        texts.append(f"SHEET: {sheet_name}")
        for _, row in df.iterrows():
            row_str = " | ".join(
                [f"{col}: {val}" for col, val in row.items() if pd.notna(val)]
            )
            if row_str:
                texts.append(row_str)

    return texts


def process_word(file_content):
    """
    Process Word file and return a list of strings representing paragraphs,
    tables, headers, and footers.
    """
    doc = Document(io.BytesIO(file_content))
    texts = []
    current_chunk = ""

    def flush_chunk():
        nonlocal current_chunk
        if current_chunk.strip():
            texts.append(current_chunk.strip())
        current_chunk = ""

    for para in doc.paragraphs:
        if _paragraph_in_table(para):
            continue

        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style is not None else ""
        if len(text) < 100 or style_name.startswith('Heading'):
            flush_chunk()
            current_chunk = text
        else:
            current_chunk += "\n" + text if current_chunk else text

        if len(current_chunk) > 1500:
            flush_chunk()

    flush_chunk()

    # Extract table content separately to preserve structure.
    for i, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            texts.append(f"TABLE {i}:")
            texts.extend(rows)

    # Extract headers and footers if present.
    for section_index, section in enumerate(doc.sections, start=1):
        header_text = "\n".join([p.text.strip() for p in section.header.paragraphs if p.text.strip()])
        footer_text = "\n".join([p.text.strip() for p in section.footer.paragraphs if p.text.strip()])
        if header_text:
            texts.append(f"HEADER (section {section_index}): {header_text}")
        if footer_text:
            texts.append(f"FOOTER (section {section_index}): {footer_text}")

    return texts


def process_pdf(file_content):
    """
    Process PDF file and return a list of strings (pages/chunks).
    """
    reader = PdfReader(io.BytesIO(file_content))
    texts = []

    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text()
        if not text or not text.strip():
            continue

        page_text = text.strip()
        if len(page_text) > 2000:
            for start in range(0, len(page_text), 2000):
                texts.append(page_text[start:start + 2000])
        else:
            texts.append(f"PAGE {page_index}: {page_text}")

    return texts


def process_text(file_content):
    """Process plain text file and return a list of text lines."""
    content = file_content.decode('utf-8', errors='ignore')
    return [line.strip() for line in content.splitlines() if line.strip()]


def process_csv(file_content):
    """Process CSV file and return a list of rows."""
    text = file_content.decode('utf-8', errors='ignore')
    reader = csv.reader(io.StringIO(text))
    texts = []
    for row in reader:
        row = [cell.strip() for cell in row if cell.strip()]
        if row:
            texts.append(' | '.join(row))
    return texts


def process_pptx(file_content):
    """Process PowerPoint file and return a list of slide texts."""
    if Presentation is None:
        raise ImportError('python-pptx is not installed. Install it to support .pptx uploads.')

    prs = Presentation(io.BytesIO(file_content))
    texts = []

    for slide_index, slide in enumerate(prs.slides, start=1):
        slide_lines = []
        for shape in slide.shapes:
            if hasattr(shape, 'text') and shape.text.strip():
                slide_lines.append(shape.text.strip())
        if slide_lines:
            texts.append(f"SLIDE {slide_index}:")
            texts.extend(slide_lines)

    return texts
